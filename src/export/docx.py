"""
DOCX exporter for legal documents.

Uses python-docx to generate Word documents with proper
court-standard formatting.
"""

import io
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT

from src.export.models import ExportFormat, ExportResult
from src.generation.generator import GeneratedDocument
from src.templates.schemas import FormattingRequirements


class DocxExporter:
    """
    Exports legal documents to DOCX with court-standard formatting.
    
    Applies formatting from template:
    - Paper size (A4)
    - Font (Times New Roman)
    - Margins (4cm left/right, 2cm top/bottom)
    - Font size (14pt)
    - Line spacing (1.5)
    
    Example:
        exporter = DocxExporter()
        result = exporter.export(generated_doc, formatting)
        docx_bytes = result.content_bytes
    """
    
    def __init__(self):
        """Initialize the DOCX exporter."""
        pass
    
    def export(
        self,
        document: GeneratedDocument,
        formatting: Optional[FormattingRequirements] = None
    ) -> ExportResult:
        """
        Export a generated document to DOCX.
        
        Args:
            document: The generated legal document
            formatting: Optional formatting requirements. If None, uses defaults.
            
        Returns:
            ExportResult with DOCX bytes and metadata
        """
        # Use default formatting if not provided
        if formatting is None:
            formatting = FormattingRequirements()
        
        # Create new document
        doc = Document()
        
        # Configure page setup
        section = doc.sections[0]
        section.page_width = Cm(21)  # A4 width
        section.page_height = Cm(29.7)  # A4 height
        section.orientation = WD_ORIENT.PORTRAIT
        
        # Parse and apply margins
        margin_lr = self._parse_margin_to_cm(formatting.margin_left_right)
        margin_tb = self._parse_margin_to_cm(formatting.margin_top_bottom)
        
        section.left_margin = Cm(margin_lr)
        section.right_margin = Cm(margin_lr)
        section.top_margin = Cm(margin_tb)
        section.bottom_margin = Cm(margin_tb)
        
        # Write content
        self._write_content(doc, document.content, formatting)
        
        # Save to bytes
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_bytes = docx_buffer.getvalue()
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{document.doc_type}_{document.court_level}_{timestamp}.docx"
        
        return ExportResult(
            format=ExportFormat.DOCX,
            content_bytes=docx_bytes,
            filename=filename,
            page_count=None  # DOCX doesn't have page count until rendered
        )
    
    def _parse_margin_to_cm(self, margin_str: str) -> float:
        """
        Parse margin string to centimeters.
        
        Args:
            margin_str: Margin string like "4cm" or "2cm"
            
        Returns:
            Margin in centimeters
        """
        margin_str = margin_str.strip().lower()
        
        if margin_str.endswith("cm"):
            return float(margin_str[:-2])
        elif margin_str.endswith("mm"):
            return float(margin_str[:-2]) / 10
        elif margin_str.endswith("in"):
            return float(margin_str[:-2]) * 2.54
        else:
            try:
                return float(margin_str)
            except ValueError:
                return 4.0  # Default 4cm
    
    def _write_content(
        self,
        doc: Document,
        content: str,
        formatting: FormattingRequirements
    ):
        """
        Write document content to DOCX with proper formatting.
        
        Args:
            doc: python-docx Document instance
            content: Document content string
            formatting: Formatting requirements
        """
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            if not para_text.strip():
                continue
            
            # Handle lines within paragraph
            lines = para_text.split('\n')
            
            for i, line in enumerate(lines):
                if not line.strip():
                    continue
                    
                # Add paragraph
                para = doc.add_paragraph()
                
                # Apply formatting
                para_format = para.paragraph_format
                
                # Line spacing (1.5 = MULTIPLE with value 1.5)
                para_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                para_format.line_spacing = formatting.line_spacing
                
                # Space after paragraph
                para_format.space_after = Pt(6)
                
                # Add text with font formatting
                run = para.add_run(line.strip())
                run.font.name = formatting.font
                run.font.size = Pt(formatting.font_size)
